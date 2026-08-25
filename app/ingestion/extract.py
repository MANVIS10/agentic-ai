"""Document text extraction, moved from stage25_react_ui/backend/main.py
(lines 829-907): validate -> extract -> chunk -> embed -> store. New in
Stage 25 (spec §3/§4): a filename length cap, a bounded read (instead of
read-then-check) for the size limit, a PDF page-count cap, a DOCX
zip-bomb guard, and a uniform extraction timeout - all four "dangerous
file" rejections collapse into the SAME existing generic 422 message,
deliberately, so none of them leaks which specific defense triggered.
"""

import asyncio
import io
import zipfile

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.config import (
    ALLOWED_FILE_TYPES,
    EXTRACTION_TIMEOUT_SECONDS,
    MAX_DOCX_UNCOMPRESSED_BYTES,
    MAX_PDF_PAGES,
)


def get_file_type(filename: str) -> str | None:
    """Return 'pdf'/'txt'/'docx' from the filename's extension, or None if
    unsupported (or the filename has no extension at all). Extension-based,
    not UploadFile.content_type - multipart clients set that inconsistently
    (many send application/octet-stream for anything), so it isn't a
    reliable signal on its own. Only the LAST extension is checked
    (rsplit), so a double-extension filename like "resume.pdf.exe" is
    already correctly rejected (extension "exe", not in ALLOWED_FILE_TYPES).
    """
    if "." not in filename:
        return None
    extension = filename.rsplit(".", 1)[-1].lower()
    return extension if extension in ALLOWED_FILE_TYPES else None


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Extract raw text from the uploaded bytes for a supported file_type.

    Raises on a corrupt/unparseable file, on a PDF over MAX_PDF_PAGES, or
    on a DOCX whose declared uncompressed size exceeds
    MAX_DOCX_UNCOMPRESSED_BYTES - the caller (upload_document) catches all
    of these identically and maps them to the same generic 422 (spec §4).
    Genuine parse failure and the two new caps are deliberately
    indistinguishable from the outside.

    This also doubles as the real, content-based check that a file's
    extension didn't lie: a .pdf-named file that isn't actually a valid
    PDF fails here, inside PdfReader, rather than being silently accepted.
    """
    if file_type == "pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError(f"PDF exceeds the maximum allowed page count ({MAX_PDF_PAGES})")
        return " ".join(
            " ".join((page.extract_text() or "").split()) for page in reader.pages
        )
    if file_type == "docx":
        # A .docx file is a ZIP archive; python-docx's Document() fully
        # decompresses and parses every entry inside it. Check the declared
        # total UNCOMPRESSED size across every entry (a standard zip-bomb
        # mitigation) BEFORE calling DocxDocument(), which would otherwise
        # do that decompression itself with no size guard. zipfile is the
        # standard library - no new dependency.
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            total_uncompressed = sum(entry.file_size for entry in archive.infolist())
        if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise ValueError(
                f"DOCX uncompressed size ({total_uncompressed} bytes) exceeds the "
                f"maximum allowed limit ({MAX_DOCX_UNCOMPRESSED_BYTES} bytes)"
            )
        document = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    # txt - strict decode (no errors="ignore") gives this the same
    # "content is the real check" property as PDF/DOCX above: a non-UTF-8
    # file raises instead of silently losing bytes.
    return file_bytes.decode("utf-8")


async def extract_text_with_timeout(file_bytes: bytes, file_type: str) -> str:
    """Runs extract_text under a hard wall-clock bound
    (EXTRACTION_TIMEOUT_SECONDS), independent of the page-count/zip-bomb
    caps above - the general-purpose safety net for parser pathologies
    neither of those anticipates (e.g. a small, low-page-count PDF with a
    deeply nested object graph that's just slow to walk).

    Phase 2 (async conversion): extract_text is CPU-bound (parsing a PDF/
    DOCX), so it must never run directly on the event loop - that would
    block every other request in the process for the whole parse, not just
    the caller. asyncio.to_thread(...) runs it on a worker thread (the
    default executor), and asyncio.wait_for(...) applies the same
    EXTRACTION_TIMEOUT_SECONDS bound the original ThreadPoolExecutor +
    future.result(timeout=...) enforced - callers must keep catching a
    timeout (now asyncio.TimeoutError, which is `TimeoutError` itself on
    this Python version, rather than concurrent.futures.TimeoutError) and
    mapping it to the same generic 422.

    The worker thread is still abandoned (not forcibly killed) on timeout -
    Python has no safe primitive to terminate a running thread, async or
    not. This is the same accepted, documented limitation as the original
    (see the stage README), not a correctness issue: the request has
    already returned its error to the caller either way, and
    asyncio.wait_for cancelling the wrapping task does not block THIS
    function waiting for the abandoned thread to actually finish.

    EXTRACTION_TIMEOUT_SECONDS and extract_text are read as free variables
    from this module's own globals on every call (not captured at import
    time), so tests can monkeypatch either name here directly.
    """
    return await asyncio.wait_for(
        asyncio.to_thread(extract_text, file_bytes, file_type),
        timeout=EXTRACTION_TIMEOUT_SECONDS,
    )
