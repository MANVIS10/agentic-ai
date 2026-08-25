"""
Stage 20: add a POST /documents/upload endpoint to Stage 19's FastAPI app,
so a user can hand the system a PDF, TXT, or DOCX file and have it
validated, text-extracted, chunked, and durably stored in PostgreSQL.

New concept vs. Stage 1-19:
- Every prior stage that touched a document (Stage 3, 5) worked with
  content the project already had access to - files bundled at build time
  (knowledge_base/*.md) or a URL fetched over HTTP (fetch_pdf). Nothing so
  far accepts arbitrary binary/text input handed to it by a caller over
  HTTP. This stage's upload_document() route does exactly that: it reads
  raw bytes from a multipart/form-data request, has to figure out what
  kind of file they are, and has to validate/extract/chunk them without
  ever trusting the caller's claims about the file (extension is checked,
  but extraction itself is the real test of "is this actually a valid
  PDF/DOCX?").
- The two new tables (documents, document_chunks) are the first tables in
  this repo created with raw hand-written SQL (CREATE TABLE IF NOT
  EXISTS), rather than being owned by a library (PostgresSaver owns
  checkpoints/checkpoint_blobs/checkpoint_writes/checkpoint_migrations).
  Same "idempotent, safe to run at every process start" convention as
  checkpointer.setup(), just written by hand instead of provided by
  langgraph-checkpoint-postgres.
- This is deliberately storage only - nothing here embeds, indexes, or
  searches the stored chunks. See .claude/spec/stage20_document_upload_spec.md
  for the full spec (Step 21 is expected to add embeddings/retrieval on
  top of the document_chunks table this stage creates, without needing to
  touch this ingestion pipeline).
- Everything else - every specialist subgraph, the supervisor+critic
  graph, the outer planner+human-approval graph, PostgresSaver
  checkpointing, per-thread_id locking, and the existing /health /chat
  /approve /reject routes - is copied verbatim from Stage 19. Nothing
  about the graph itself changed.

See this folder's README.md for how to start Postgres, run the API, and
exercise the new endpoint.
"""

import ast
import io
import operator
import os
import threading
import uuid
from contextlib import contextmanager
from pathlib import Path
from typing import Literal

import psycopg
import uvicorn
from docx import Document as DocxDocument
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.tools import tool
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langgraph.checkpoint.postgres import PostgresSaver
from langgraph.graph import END, START, MessagesState, StateGraph
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.types import Command, interrupt
from pydantic import BaseModel
from pypdf import PdfReader
from typing_extensions import TypedDict

load_dotenv()

DATABASE_URL = os.getenv(
    "DATABASE_URL", "postgresql://postgres:postgres@localhost:5433/postgres?sslmode=disable"
)

MAX_RETRIES = 1  # at most one retry per subtask - two specialist attempts total

llm = ChatOpenAI(model="gpt-4o-mini")


# ---------------------------------------------------------------------------
# Research Agent - web search specialist (copied from Stage 19, unchanged)
# ---------------------------------------------------------------------------

RESEARCH_SYSTEM_PROMPT = (
    "You are a Research Agent, a specialist whose only job is web research. "
    "You have one tool: web search. When asked something you don't already "
    "know for certain, search the web for it before answering. Report what "
    "you found clearly and cite that it came from a web search when you use "
    "it. Stay focused on research - you're not a general-purpose assistant."
)

search_web = DuckDuckGoSearchRun()
research_tools = [search_web]

research_llm = ChatOpenAI(model="gpt-4o-mini")
research_llm_with_tools = research_llm.bind_tools(research_tools)


def research_agent_node(state: MessagesState):
    messages = [SystemMessage(content=RESEARCH_SYSTEM_PROMPT), *state["messages"]]
    response = research_llm_with_tools.invoke(messages)
    return {"messages": [response]}


research_subgraph_builder = StateGraph(MessagesState)
research_subgraph_builder.add_node("agent", research_agent_node)
research_subgraph_builder.add_node("tools", ToolNode(research_tools))
research_subgraph_builder.add_edge(START, "agent")
research_subgraph_builder.add_conditional_edges("agent", tools_condition)
research_subgraph_builder.add_edge("tools", "agent")

research_graph = research_subgraph_builder.compile()


# ---------------------------------------------------------------------------
# Knowledge Agent - local knowledge-base specialist (copied from Stage 19,
# unchanged, using this folder's own knowledge_base/ copy)
# ---------------------------------------------------------------------------

KNOWLEDGE_SYSTEM_PROMPT = (
    "You are a Knowledge Agent, a specialist whose only job is answering "
    "questions from a local knowledge base of documents. You have one tool: "
    "knowledge-base search. You cannot browse the web or access anything "
    "outside these documents. If the knowledge base doesn't contain the "
    "answer, say so plainly instead of guessing. Stay focused on the "
    "knowledge base - you're not a general-purpose assistant."
)

KNOWLEDGE_BASE_DIR = Path(__file__).parent / "knowledge_base"


def load_knowledge_base() -> list[Document]:
    """Read every .md file in knowledge_base/ and split it into chunks."""
    splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=50)
    documents = []
    for path in sorted(KNOWLEDGE_BASE_DIR.glob("*.md")):
        text = path.read_text(encoding="utf-8")
        for chunk in splitter.split_text(text):
            documents.append(Document(page_content=chunk, metadata={"source": path.name}))
    return documents


embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
vector_store = InMemoryVectorStore(embeddings)
vector_store.add_documents(load_knowledge_base())


@tool
def search_knowledge_base(query: str) -> str:
    """Search the local knowledge base for information relevant to a
    natural-language question and return the most relevant text chunks.

    Use this when the user asks something that could be answered from the
    project's own documents (currently: renewable energy topics - solar,
    wind, and hydro power) rather than general knowledge or current events.
    """
    results = vector_store.similarity_search(query, k=3)
    if not results:
        return "No relevant information found in the knowledge base."

    formatted = []
    for doc in results:
        source = doc.metadata.get("source", "unknown")
        formatted.append(f"[source: {source}]\n{doc.page_content}")
    return "\n\n".join(formatted)


knowledge_tools = [search_knowledge_base]

knowledge_llm = ChatOpenAI(model="gpt-4o-mini")
knowledge_llm_with_tools = knowledge_llm.bind_tools(knowledge_tools)


def knowledge_agent_node(state: MessagesState):
    messages = [SystemMessage(content=KNOWLEDGE_SYSTEM_PROMPT), *state["messages"]]
    response = knowledge_llm_with_tools.invoke(messages)
    return {"messages": [response]}


knowledge_subgraph_builder = StateGraph(MessagesState)
knowledge_subgraph_builder.add_node("agent", knowledge_agent_node)
knowledge_subgraph_builder.add_node("tools", ToolNode(knowledge_tools))
knowledge_subgraph_builder.add_edge(START, "agent")
knowledge_subgraph_builder.add_conditional_edges("agent", tools_condition)
knowledge_subgraph_builder.add_edge("tools", "agent")

knowledge_graph = knowledge_subgraph_builder.compile()


# ---------------------------------------------------------------------------
# Analysis Agent - safe-arithmetic specialist (copied from Stage 19,
# unchanged)
# ---------------------------------------------------------------------------

_ALLOWED_BINOPS = {
    ast.Add: operator.add,
    ast.Sub: operator.sub,
    ast.Mult: operator.mul,
    ast.Div: operator.truediv,
    ast.FloorDiv: operator.floordiv,
    ast.Mod: operator.mod,
    ast.Pow: operator.pow,
}

_ALLOWED_UNARYOPS = {
    ast.UAdd: operator.pos,
    ast.USub: operator.neg,
}


def _eval_node(node):
    if isinstance(node, ast.Expression):
        return _eval_node(node.body)
    if isinstance(node, ast.Constant):
        if isinstance(node.value, (int, float)):
            return node.value
        raise ValueError(f"Unsupported constant: {node.value!r}")
    if isinstance(node, ast.BinOp) and type(node.op) in _ALLOWED_BINOPS:
        left = _eval_node(node.left)
        right = _eval_node(node.right)
        return _ALLOWED_BINOPS[type(node.op)](left, right)
    if isinstance(node, ast.UnaryOp) and type(node.op) in _ALLOWED_UNARYOPS:
        return _ALLOWED_UNARYOPS[type(node.op)](_eval_node(node.operand))
    raise ValueError(f"Unsupported expression: {ast.dump(node)}")


@tool
def calculate(expression: str) -> str:
    """Evaluate a numeric arithmetic expression and return the result.

    Supports +, -, *, /, // (floor division), % (modulo), ** (power), and
    parentheses. Use this for anything requiring exact arithmetic - sums,
    averages (e.g. "(12 + 18 + 30) / 3"), percentage change
    (e.g. "((120 - 100) / 100) * 100"), or differences between values -
    instead of computing it yourself, since arithmetic done by an LLM is
    unreliable.
    """
    try:
        tree = ast.parse(expression, mode="eval")
        result = _eval_node(tree)
    except Exception as exc:
        return f"Could not evaluate '{expression}': {exc}"
    return str(result)


analysis_tools = [calculate]

ANALYSIS_SYSTEM_PROMPT = (
    "You are an Analysis Agent, a specialist in calculations, comparisons, "
    "and reasoning over numeric or structured data. You work only with "
    "numbers and data given to you in the conversation - you cannot search "
    "the web or read documents. You have one tool: calculate, which "
    "evaluates arithmetic expressions exactly. Use it for any sum, average, "
    "percentage change, or difference instead of computing by hand, since "
    "you are prone to arithmetic mistakes. For comparisons like 'which is "
    "highest', reason directly over the numbers once you have them. Stay "
    "focused on analysis - you're not a general-purpose assistant."
)

analysis_llm = ChatOpenAI(model="gpt-4o-mini")
analysis_llm_with_tools = analysis_llm.bind_tools(analysis_tools)


def analysis_agent_node(state: MessagesState):
    messages = [SystemMessage(content=ANALYSIS_SYSTEM_PROMPT), *state["messages"]]
    response = analysis_llm_with_tools.invoke(messages)
    return {"messages": [response]}


analysis_subgraph_builder = StateGraph(MessagesState)
analysis_subgraph_builder.add_node("agent", analysis_agent_node)
analysis_subgraph_builder.add_node("tools", ToolNode(analysis_tools))
analysis_subgraph_builder.add_edge(START, "agent")
analysis_subgraph_builder.add_conditional_edges("agent", tools_condition)
analysis_subgraph_builder.add_edge("tools", "agent")

analysis_graph = analysis_subgraph_builder.compile()


# ---------------------------------------------------------------------------
# Supervisor + Critic - routing and review (copied from Stage 19, unchanged)
# ---------------------------------------------------------------------------


class CriticState(MessagesState):
    next: Literal["research", "knowledge", "analysis"]
    verdict: Literal["pass", "retry"]
    feedback: str
    retry_count: int


class Route(TypedDict):
    """The supervisor's routing decision."""

    next: Literal["research", "knowledge", "analysis"]


SUPERVISOR_SYSTEM_PROMPT = (
    "You are a supervisor that routes a user's question to exactly one of "
    "three specialist agents:\n\n"
    "- 'research': a Research Agent that searches the live web. Use this "
    "for current events, recent news, or anything that changes over time "
    "and needs up-to-date information.\n"
    "- 'knowledge': a Knowledge Agent that searches a local knowledge base "
    "covering renewable energy topics (solar, wind, and hydro power). Use "
    "this for questions about those topics.\n"
    "- 'analysis': an Analysis Agent that performs exact arithmetic - sums, "
    "averages, percentage changes, and comparisons - over numbers already "
    "given in the conversation. Use this for calculation questions; it "
    "cannot search the web or read documents.\n\n"
    "Read the user's latest question and decide which one specialist "
    "should handle it."
)

# method="function_calling" (a real tool call) instead of the default
# "json_schema": gpt-4o-mini's default structured-output mode occasionally
# echoes back the JSON SCHEMA itself instead of an instance of it, which
# crashes a dict lookup with a KeyError. function_calling reliably returns
# just the typed fields.
supervisor_llm = ChatOpenAI(model="gpt-4o-mini").with_structured_output(
    Route, method="function_calling"
)


def supervisor_node(state: CriticState):
    messages = [SystemMessage(content=SUPERVISOR_SYSTEM_PROMPT), *state["messages"]]
    route: Route = supervisor_llm.invoke(messages)
    return {"next": route["next"], "retry_count": 0}


def research_node(state: CriticState):
    """Run the Research Agent subgraph and hand back only its final answer."""
    messages = state["messages"]
    if state.get("feedback"):
        messages = messages + [
            HumanMessage(
                content=f"Reviewer feedback: {state['feedback']} "
                "Please address this and try again."
            )
        ]
    result = research_graph.invoke({"messages": messages})
    return {"messages": [result["messages"][-1]]}


def knowledge_node(state: CriticState):
    """Run the Knowledge Agent subgraph and hand back only its final answer."""
    messages = state["messages"]
    if state.get("feedback"):
        messages = messages + [
            HumanMessage(
                content=f"Reviewer feedback: {state['feedback']} "
                "Please address this and try again."
            )
        ]
    result = knowledge_graph.invoke({"messages": messages})
    return {"messages": [result["messages"][-1]]}


def analysis_node(state: CriticState):
    """Run the Analysis Agent subgraph and hand back only its final answer."""
    messages = state["messages"]
    if state.get("feedback"):
        messages = messages + [
            HumanMessage(
                content=f"Reviewer feedback: {state['feedback']} "
                "Please address this and try again."
            )
        ]
    result = analysis_graph.invoke({"messages": messages})
    return {"messages": [result["messages"][-1]]}


def route_from_supervisor(state: CriticState) -> str:
    return state["next"]


class Review(TypedDict):
    """The critic's judgment of a specialist's answer."""

    verdict: Literal["pass", "retry"]
    feedback: str


CRITIC_SYSTEM_PROMPT = (
    "You are a critic reviewing whether an answer adequately addresses a "
    "user's question. Pass if the answer is relevant, reasonably complete, "
    "and not a refusal or empty non-answer. Retry if it clearly misses the "
    "question, is empty, or is far too vague to be useful. If you retry, "
    "give one short sentence of feedback describing what's wrong; leave "
    "feedback empty if you pass."
)

critic_llm = ChatOpenAI(model="gpt-4o-mini").with_structured_output(
    Review, method="function_calling"
)


def critic_node(state: CriticState):
    question = state["messages"][0].content
    answer = state["messages"][-1].content
    messages = [
        SystemMessage(content=CRITIC_SYSTEM_PROMPT),
        HumanMessage(content=f"Question: {question}\n\nAnswer: {answer}"),
    ]
    review: Review = critic_llm.invoke(messages)

    if review["verdict"] == "retry" and state["retry_count"] < MAX_RETRIES:
        return {
            "verdict": "retry",
            "feedback": review["feedback"],
            "retry_count": state["retry_count"] + 1,
        }
    return {"verdict": "pass", "feedback": ""}


def route_from_critic(state: CriticState) -> str:
    if state["verdict"] == "pass":
        return "end"
    return state["next"]


supervisor_critic_builder = StateGraph(CriticState)
supervisor_critic_builder.add_node("supervisor", supervisor_node)
supervisor_critic_builder.add_node("research_agent", research_node)
supervisor_critic_builder.add_node("knowledge_agent", knowledge_node)
supervisor_critic_builder.add_node("analysis_agent", analysis_node)
supervisor_critic_builder.add_node("critic", critic_node)

supervisor_critic_builder.add_edge(START, "supervisor")
supervisor_critic_builder.add_conditional_edges(
    "supervisor",
    route_from_supervisor,
    {
        "research": "research_agent",
        "knowledge": "knowledge_agent",
        "analysis": "analysis_agent",
    },
)
supervisor_critic_builder.add_edge("research_agent", "critic")
supervisor_critic_builder.add_edge("knowledge_agent", "critic")
supervisor_critic_builder.add_edge("analysis_agent", "critic")
supervisor_critic_builder.add_conditional_edges(
    "critic",
    route_from_critic,
    {
        "research": "research_agent",
        "knowledge": "knowledge_agent",
        "analysis": "analysis_agent",
        "end": END,
    },
)

# No checkpointer here on purpose: this graph is invoked as a one-shot
# helper once per subtask (like Stage 8's research_agent), not run as its
# own multi-turn REPL the way Stage 16 runs it. Only the OUTER planner graph
# below needs a checkpointer, for its one interrupt().
supervisor_critic_graph = supervisor_critic_builder.compile()


# ---------------------------------------------------------------------------
# Planner + human approval (copied from Stage 19, unchanged)
# ---------------------------------------------------------------------------


class PlannerState(TypedDict):
    question: str
    subtasks: list[str]
    current_index: int
    results: list[str]
    final_answer: str
    approved: bool


def plan(state: PlannerState):
    prompt = (
        "Break the following research question into 2-3 short, concrete "
        "subtasks that could each be researched independently. "
        "Reply with just the subtasks, one per line, no numbering.\n\n"
        f"Question: {state['question']}"
    )
    response = llm.invoke(prompt)
    subtasks = [line.strip() for line in response.content.splitlines() if line.strip()]

    print(f"\nPlan ({len(subtasks)} subtasks):")
    for i, subtask in enumerate(subtasks, start=1):
        print(f"  {i}. {subtask}")

    # Reset every field a later node might set, not just the ones this node
    # itself uses. plan() is the one node guaranteed to run first on every
    # turn (START -> plan), so it's the only place that can undo a stale
    # final_answer/approved left behind by an earlier question on the same
    # thread.
    return {
        "subtasks": subtasks,
        "current_index": 0,
        "results": [],
        "final_answer": "",
        "approved": False,
    }


def human_approval(state: PlannerState):
    decision = interrupt("Approve this plan? (y/n): ")
    return {"approved": decision.strip().lower() == "y"}


def route_after_approval(state: PlannerState) -> str:
    if not state["approved"]:
        return END
    # Don't assume plan() produced at least one subtask - reuse the same
    # "anything left to research?" check used between subtask loops, so an
    # empty plan goes straight to synthesize instead of indexing into an
    # empty subtasks list in research_subtask.
    return has_more_subtasks(state)


def research_subtask(state: PlannerState):
    """Research one subtask by running it through the full
    supervisor -> specialist -> critic pipeline, instead of a bare LLM call
    (Stage 6/7) or a single flat tool agent (Stage 8).

    Each subtask gets a fresh invocation - no shared thread/state carries
    over between subtasks, so retry_count always starts at 0 for each one.
    """
    subtask = state["subtasks"][state["current_index"]]
    print(f"\nResearching: {subtask}")

    result = supervisor_critic_graph.invoke(
        {"messages": [{"role": "user", "content": subtask}]}
    )
    print(f"  [Supervisor routed to: {result['next']}]")
    print(f"  [Critic verdict: {result['verdict']}, retries used: {result['retry_count']}]")

    answer = result["messages"][-1].content

    return {
        "results": state["results"] + [answer],
        "current_index": state["current_index"] + 1,
    }


def has_more_subtasks(state: PlannerState) -> str:
    if state["current_index"] < len(state["subtasks"]):
        return "research_subtask"
    return "synthesize"


def synthesize(state: PlannerState):
    subtasks_and_results = "\n\n".join(
        f"Subtask: {subtask}\nAnswer: {result}"
        for subtask, result in zip(state["subtasks"], state["results"])
    )
    prompt = (
        f"Original question: {state['question']}\n\n"
        f"Research notes:\n{subtasks_and_results}\n\n"
        "Combine these into one clear final answer to the original question."
    )
    response = llm.invoke(prompt)
    return {"final_answer": response.content}


graph_builder = StateGraph(PlannerState)
graph_builder.add_node("plan", plan)
graph_builder.add_node("human_approval", human_approval)
graph_builder.add_node("research_subtask", research_subtask)
graph_builder.add_node("synthesize", synthesize)

graph_builder.add_edge(START, "plan")
graph_builder.add_edge("plan", "human_approval")
graph_builder.add_conditional_edges("human_approval", route_after_approval)
graph_builder.add_conditional_edges("research_subtask", has_more_subtasks)
graph_builder.add_edge("synthesize", END)

# Same checkpointer setup as Stage 19: a PostgresSaver backed by a real
# database, so this graph's checkpoints (the plan, subtasks, results-so-far,
# and any paused-at-human_approval interrupt) outlive the Python process.
# The connection is opened once at module scope so `graph` (and
# `checkpointer`) exist as importable, already-connected module-level names
# wherever this module is imported from - by uvicorn when it loads `app`,
# or by this stage's own test script (`from main import app, pg_conn`).
pg_conn = psycopg.connect(DATABASE_URL, autocommit=True, prepare_threshold=0)
checkpointer = PostgresSaver(pg_conn)
checkpointer.setup()  # idempotent: creates the checkpoint tables on first run only
graph = graph_builder.compile(checkpointer=checkpointer)


# ---------------------------------------------------------------------------
# Document upload tables (Stage 20 - new). Created idempotently at module
# load, the same "safe to call every process start" convention as
# checkpointer.setup() just above - but hand-written SQL, since these two
# tables aren't owned by any library the way the checkpoint tables are
# owned by langgraph-checkpoint-postgres. No migrations framework, per this
# project's minimal-dependencies rule.
# ---------------------------------------------------------------------------

DOCUMENTS_TABLE_SQL = """
CREATE TABLE IF NOT EXISTS documents (
    id UUID PRIMARY KEY,
    filename TEXT NOT NULL,
    file_type TEXT NOT NULL,
    file_size_bytes INTEGER NOT NULL,
    chunk_count INTEGER NOT NULL,
    uploaded_at TIMESTAMPTZ NOT NULL DEFAULT now()
)
"""

DOCUMENT_CHUNKS_TABLE_SQL = """
CREATE TABLE IF NOT EXISTS document_chunks (
    id UUID PRIMARY KEY,
    document_id UUID NOT NULL REFERENCES documents(id) ON DELETE CASCADE,
    chunk_index INTEGER NOT NULL,
    content TEXT NOT NULL,
    created_at TIMESTAMPTZ NOT NULL DEFAULT now()
)
"""

pg_conn.execute(DOCUMENTS_TABLE_SQL)
pg_conn.execute(DOCUMENT_CHUNKS_TABLE_SQL)  # after documents - it has a FK reference to it


# ---------------------------------------------------------------------------
# Document upload pipeline (Stage 20 - new): validate -> extract -> chunk.
# Storage happens inline in the upload_document() route below, since it
# needs the document_id generated there.
# ---------------------------------------------------------------------------

ALLOWED_FILE_TYPES = {"pdf", "txt", "docx"}
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB

# Same chunk_size/chunk_overlap as load_knowledge_base() above, for
# consistency across every place this project chunks text.
CHUNK_SIZE = 400
CHUNK_OVERLAP = 50
document_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
)


def get_file_type(filename: str) -> str | None:
    """Return 'pdf'/'txt'/'docx' from the filename's extension, or None if
    unsupported (or the filename has no extension at all). Extension-based,
    not UploadFile.content_type - multipart clients set that inconsistently
    (many send application/octet-stream for anything), so it isn't a
    reliable signal on its own.
    """
    if "." not in filename:
        return None
    extension = filename.rsplit(".", 1)[-1].lower()
    return extension if extension in ALLOWED_FILE_TYPES else None


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Extract raw text from the uploaded bytes for a supported file_type.

    Raises on a corrupt/unparseable file - the caller (upload_document)
    catches that and maps it to a 422. This doubles as the real, content-
    based check that a file's extension didn't lie: a .pdf-named file that
    isn't actually a valid PDF fails here, inside PdfReader, rather than
    being silently accepted.
    """
    if file_type == "pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        return " ".join(
            " ".join((page.extract_text() or "").split()) for page in reader.pages
        )
    if file_type == "docx":
        document = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    # txt - strict decode (no errors="ignore") gives this the same
    # "content is the real check" property as PDF/DOCX above: a non-UTF-8
    # file raises instead of silently losing bytes.
    return file_bytes.decode("utf-8")


# ---------------------------------------------------------------------------
# Per-thread_id locking - an in-process guard in front of /chat, /approve,
# and /reject, closing a race that has nothing to do with the graph or
# Postgres being wrong: graph.invoke({"question": ...}, config) in /chat
# takes real wall-clock time (plan()'s LLM call), and doesn't commit its
# "paused at human_approval" checkpoint until it returns. If /approve or
# /reject for the SAME thread_id runs graph.get_state(config) before that
# commit lands - e.g. a client that doesn't wait for /chat's response
# before firing the next request - it reads whatever checkpoint was
# PREVIOUSLY the latest for that thread_id (nothing, for a brand-new
# thread_id -> 404; a prior completed run, for a reused one -> a stale,
# misleading 409 "not currently awaiting approval").
#
# The fix isn't in the checkpointer or the graph - both are already correct
# and consistent, just read at the wrong moment. It's mutual exclusion
# around "read/act on this thread_id's state", scoped to one Python process
# (this stage runs as a single sync process - see the README's "Design
# decisions" for why that's the right scope here, and what would need to
# change for a multi-worker deployment). Copied from Stage 19, unchanged -
# the new /documents/upload route below doesn't use it, since uploads
# aren't scoped to a thread_id.
# ---------------------------------------------------------------------------

# How long a request will wait for another request already in flight on the
# same thread_id before giving up. Generous on purpose: /approve can hold
# this for the entire research_subtask loop (multiple specialist + critic
# calls across 2-3 subtasks, with possible retries) - not just the few
# seconds /chat's plan() call takes.
THREAD_LOCK_TIMEOUT_SECONDS = 120

_thread_locks: dict[str, threading.Lock] = {}
_thread_locks_guard = threading.Lock()  # protects _thread_locks itself, not held while a per-thread lock is held


def _lock_for(thread_id: str) -> threading.Lock:
    """One Lock per thread_id, created on first use. Different thread_ids
    never block each other - only two requests for the SAME thread_id
    contend.
    """
    with _thread_locks_guard:
        if thread_id not in _thread_locks:
            _thread_locks[thread_id] = threading.Lock()
        return _thread_locks[thread_id]


@contextmanager
def _thread_lock(thread_id: str):
    """Hold this thread_id's lock for the life of one request, so a
    concurrent /chat, /approve, or /reject for the SAME thread_id can never
    read Postgres state while this request is still in the middle of
    changing it. Raises the same 409 /approve and /reject already use for
    "not currently awaiting approval" if the other request hasn't finished
    within THREAD_LOCK_TIMEOUT_SECONDS, rather than blocking forever.
    """
    lock = _lock_for(thread_id)
    if not lock.acquire(timeout=THREAD_LOCK_TIMEOUT_SECONDS):
        raise HTTPException(
            status_code=409,
            detail="This thread is busy processing another request. Please try again shortly.",
        )
    try:
        yield
    finally:
        lock.release()


# ---------------------------------------------------------------------------
# FastAPI layer. /health, /chat, /approve, /reject are copied from Stage 19,
# unchanged. POST /documents/upload (Stage 20 - new) is appended after them.
# ---------------------------------------------------------------------------


class ChatRequest(BaseModel):
    question: str
    thread_id: str


class ApproveRequest(BaseModel):
    thread_id: str


class RejectRequest(BaseModel):
    thread_id: str


class ThreadStatusResponse(BaseModel):
    """Shared response shape for /chat, /approve, and /reject - which
    optional fields are populated depends on `status`:
      - "awaiting_approval" (only ever returned by /chat, given this
        graph's fixed shape): subtasks + approval_prompt are set.
      - "completed" (returned by /approve): subtasks, results, and
        final_answer are all set.
      - "rejected" (returned by /reject): subtasks is set (the plan that
        was declined); results is []; final_answer is "".
    """

    thread_id: str
    status: Literal["awaiting_approval", "completed", "rejected"]
    subtasks: list[str] | None = None
    approval_prompt: str | None = None
    results: list[str] | None = None
    final_answer: str | None = None


class HealthResponse(BaseModel):
    status: Literal["ok"]
    database: Literal["connected"]


class UploadResponse(BaseModel):
    document_id: str
    filename: str
    file_type: str
    chunk_count: int
    status: Literal["stored"]


app = FastAPI(
    title="Stage 20: Multi-Agent Research Assistant API with Document Upload",
    description=(
        "FastAPI wrapper around Stage 19's planner + human-approval + "
        "supervisor/critic/specialist research graph, checkpointed to "
        "PostgreSQL, plus a POST /documents/upload endpoint that validates, "
        "extracts, chunks, and durably stores user-uploaded PDF/TXT/DOCX "
        "files."
    ),
)


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    """Defense-in-depth net for anything that escapes a route's own
    try/except below (e.g. a failure while parsing the request itself).
    Never echoes exc's text back to the client - only logs it server-side.
    """
    print(f"[unhandled] {request.method} {request.url.path}: {exc}")
    return JSONResponse(status_code=500, content={"detail": "An unexpected error occurred."})


@app.get("/health", response_model=HealthResponse)
def health():
    """Verify the API process is up AND its Postgres dependency is
    reachable - a plain "the process is running" check would pass even if
    the database (the thing every other endpoint actually depends on) were
    down.
    """
    try:
        pg_conn.execute("SELECT 1").fetchone()
    except Exception as exc:
        print(f"[/health] Database unavailable: {exc}")
        raise HTTPException(status_code=503, detail="Database unavailable")
    return HealthResponse(status="ok", database="connected")


@app.post("/chat", response_model=ThreadStatusResponse)
def chat(request: ChatRequest):
    """Start (or restart) a research question on the given thread_id.

    Because human_approval() unconditionally calls interrupt(), this always
    pauses there and returns - it never runs research_subtask/synthesize in
    the same call. Approval/rejection happens via separate requests below.

    Held under this thread_id's lock for the whole call: plan()'s LLM call
    takes real wall-clock time, and a concurrent /approve or /reject for the
    same thread_id must not be able to read Postgres until the checkpoint
    this call is about to write has actually committed.
    """
    with _thread_lock(request.thread_id):
        config = {"configurable": {"thread_id": request.thread_id}}

        try:
            result = graph.invoke({"question": request.question}, config=config)
        except Exception as exc:
            print(f"[/chat] Error for thread_id={request.thread_id!r}: {exc}")
            raise HTTPException(
                status_code=500,
                detail="Something went wrong processing this question. Please try again.",
            )

        if "__interrupt__" in result:
            prompt = result["__interrupt__"][0].value
            return ThreadStatusResponse(
                thread_id=request.thread_id,
                status="awaiting_approval",
                subtasks=result.get("subtasks", []),
                approval_prompt=prompt,
            )

        # Shouldn't happen given the current graph (human_approval always
        # interrupts) - kept as a safety net rather than assuming the shape
        # above is the only possible outcome.
        return ThreadStatusResponse(
            thread_id=request.thread_id,
            status="completed",
            subtasks=result.get("subtasks", []),
            results=result.get("results", []),
            final_answer=result.get("final_answer", ""),
        )


def _require_pending_approval(thread_id: str):
    """Shared validation for /approve and /reject: confirm this thread_id
    actually exists and is currently paused at human_approval before
    calling Command(resume=...) on it. graph.invoke()'s return value alone
    can't answer this anymore, since the approval decision now arrives as
    its own separate request instead of the same call that produced the
    interrupt.
    """
    config = {"configurable": {"thread_id": thread_id}}
    state = graph.get_state(config)

    if not state.values:
        raise HTTPException(
            status_code=404, detail="No conversation found for this thread_id"
        )
    if "human_approval" not in state.next:
        raise HTTPException(
            status_code=409, detail="This thread is not currently awaiting approval"
        )
    return config


@app.post("/approve", response_model=ThreadStatusResponse)
def approve(request: ApproveRequest):
    """Resume a paused thread with an approval, running
    research_subtask (looped over every subtask) -> synthesize -> END.

    Held under this thread_id's lock for the whole call, so the pending-
    approval check and the resume happen atomically together: a concurrent
    /chat for the same thread_id can't slip a new checkpoint in between the
    check and the resume, and a duplicate simultaneous /approve is forced to
    wait and then see this call's result rather than racing it.
    """
    with _thread_lock(request.thread_id):
        config = _require_pending_approval(request.thread_id)

        try:
            result = graph.invoke(Command(resume="y"), config=config)
        except Exception as exc:
            print(f"[/approve] Error for thread_id={request.thread_id!r}: {exc}")
            raise HTTPException(
                status_code=500,
                detail="Something went wrong while processing the approved plan. Please try again.",
            )

        return ThreadStatusResponse(
            thread_id=request.thread_id,
            status="completed",
            subtasks=result.get("subtasks", []),
            results=result.get("results", []),
            final_answer=result.get("final_answer", ""),
        )


@app.post("/reject", response_model=ThreadStatusResponse)
def reject(request: RejectRequest):
    """Resume a paused thread with a rejection. route_after_approval sends
    this straight to END - no special-case handling needed here, results
    stays [] and final_answer stays "" since no research ever ran.

    Held under this thread_id's lock for the same reason /approve is - see
    its docstring.
    """
    with _thread_lock(request.thread_id):
        config = _require_pending_approval(request.thread_id)

        try:
            result = graph.invoke(Command(resume="n"), config=config)
        except Exception as exc:
            print(f"[/reject] Error for thread_id={request.thread_id!r}: {exc}")
            raise HTTPException(
                status_code=500,
                detail="Something went wrong while processing the rejection. Please try again.",
            )

        return ThreadStatusResponse(
            thread_id=request.thread_id,
            status="rejected",
            subtasks=result.get("subtasks", []),
            results=result.get("results", []),
            final_answer=result.get("final_answer", ""),
        )


UNSUPPORTED_TYPE_DETAIL = "Unsupported file type. Allowed types: pdf, txt, docx"


@app.post("/documents/upload", response_model=UploadResponse)
def upload_document(file: UploadFile = File(...)):
    """Validate, extract, chunk, and durably store an uploaded PDF/TXT/DOCX
    file. Storage only - nothing here makes the stored chunks searchable
    (see the module docstring and .claude/spec/stage20_document_upload_spec.md).

    Validation order matches the spec exactly - cheapest/no-I/O checks
    first: extension -> read bytes -> empty -> size limit -> extract
    (content-based check #2) -> empty-extracted-text -> chunk -> store.
    Error-handling style matches /chat, /approve, /reject: a short,
    hand-written detail string on every HTTPException, the real exception
    printed server-side, never echoed to the client.
    """
    filename = file.filename or ""
    file_type = get_file_type(filename)
    if file_type is None:
        raise HTTPException(status_code=400, detail=UNSUPPORTED_TYPE_DETAIL)

    # Sync read, same as every other route here being a plain `def` -
    # FastAPI/Starlette run sync routes in a threadpool automatically, so
    # this blocking call doesn't need to be async.
    file_bytes = file.file.read()

    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File exceeds the maximum allowed size of {MAX_FILE_SIZE_BYTES // (1024 * 1024)} MB",
        )

    try:
        text = extract_text(file_bytes, file_type)
    except Exception as exc:
        print(f"[/documents/upload] Extraction failed for {filename!r}: {exc}")
        raise HTTPException(
            status_code=422,
            detail="Could not read this file — it may be corrupted or malformed",
        )

    if not text.strip():
        raise HTTPException(
            status_code=422, detail="No extractable text found in this document"
        )

    chunks = document_splitter.split_text(text)
    document_id = uuid.uuid4()

    try:
        # pg_conn is autocommit; .transaction() still gives an explicit
        # BEGIN/COMMIT/ROLLBACK block on it, so a failure partway through
        # chunk insertion rolls back the whole thing - no orphaned
        # documents row with a wrong chunk_count or a partial chunk set.
        with pg_conn.transaction():
            pg_conn.execute(
                "INSERT INTO documents (id, filename, file_type, file_size_bytes, chunk_count) "
                "VALUES (%s, %s, %s, %s, %s)",
                (document_id, filename, file_type, len(file_bytes), len(chunks)),
            )
            for index, chunk_text in enumerate(chunks):
                pg_conn.execute(
                    "INSERT INTO document_chunks (id, document_id, chunk_index, content) "
                    "VALUES (%s, %s, %s, %s)",
                    (uuid.uuid4(), document_id, index, chunk_text),
                )
    except Exception as exc:
        print(f"[/documents/upload] DB write failed for {filename!r}: {exc}")
        raise HTTPException(
            status_code=500,
            detail="Something went wrong while storing this document. Please try again.",
        )

    return UploadResponse(
        document_id=str(document_id),
        filename=filename,
        file_type=file_type,
        chunk_count=len(chunks),
        status="stored",
    )


if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8000)
