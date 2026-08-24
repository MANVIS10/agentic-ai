"""
Smoke test for Stage 20's POST /documents/upload - validates, extracts,
chunks, and durably stores PDF/TXT/DOCX uploads in Postgres.

Not a pytest suite - same standalone-script convention as every other
stage's test file (asserts + prints, run directly with `python`). Uses
FastAPI's TestClient, so no running server is needed.

Unlike Stage 18/19's checkpoint-table tests, documents/document_chunks
aren't behind a checkpointer abstraction (there's no delete_thread()-style
helper for them), so this file queries and cleans them up directly with
the same `pg_conn` main.py already constructs - the first place in this
repo asserting against custom tables via raw psycopg.

Run with:
    python stage20_document_upload/test_document_upload.py
"""

import io

import requests
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from main import CHUNK_SIZE, MAX_FILE_SIZE_BYTES, app, pg_conn

client = TestClient(app)

# Same fixture URL stage5_pdf_fetch/test_fetch_pdf.py already uses to fetch
# a real PDF at test time - reused here instead of hand-rolling PDF bytes
# or committing a binary fixture file.
PDF_URL = "https://www.ijtsrd.com/papers/ijtsrd49820.pdf"

TEST_FILENAMES = [
    "test-report.pdf",
    "test-notes.txt",
    "test-memo.docx",
    "test-empty.txt",
    "test-virus.exe",
    "test-corrupt.pdf",
    "test-corrupt.docx",
    "test-blank.txt",
    "test-long.txt",
    "test-toobig.txt",
]


def clear_previous_test_documents():
    """Delete rows this test's fixed filenames left behind in a previous
    run (document_chunks cascade-deletes automatically via ON DELETE
    CASCADE). Adapts Stage 18/19's clear_previous_test_threads()/
    delete_thread() cleanup pattern to raw SQL, since there's no equivalent
    helper for these tables.
    """
    pg_conn.execute("DELETE FROM documents WHERE filename = ANY(%s)", (TEST_FILENAMES,))


def upload(filename, content_bytes, content_type):
    return client.post(
        "/documents/upload",
        files={"file": (filename, content_bytes, content_type)},
    )


def get_document_row(document_id):
    return pg_conn.execute(
        "SELECT filename, file_type, chunk_count FROM documents WHERE id = %s",
        (document_id,),
    ).fetchone()


def get_document_row_by_filename(filename):
    return pg_conn.execute(
        "SELECT id FROM documents WHERE filename = %s", (filename,)
    ).fetchone()


def get_chunk_rows(document_id):
    return pg_conn.execute(
        "SELECT chunk_index, content FROM document_chunks WHERE document_id = %s "
        "ORDER BY chunk_index",
        (document_id,),
    ).fetchall()


def assert_stored_and_matches_db(body):
    row = get_document_row(body["document_id"])
    assert row is not None, f"Expected a documents row for {body['document_id']}"
    assert row[0] == body["filename"], f"filename mismatch: {row[0]!r} != {body['filename']!r}"
    assert row[1] == body["file_type"], f"file_type mismatch: {row[1]!r} != {body['file_type']!r}"
    assert row[2] == body["chunk_count"], f"chunk_count mismatch: {row[2]} != {body['chunk_count']}"

    chunks = get_chunk_rows(body["document_id"])
    assert len(chunks) == body["chunk_count"], (
        f"Expected {body['chunk_count']} document_chunks rows, found {len(chunks)}"
    )
    assert [c[0] for c in chunks] == list(range(len(chunks))), (
        "chunk_index values must be contiguous 0..N-1 and in order"
    )


def run_valid_pdf_upload_check():
    pdf_bytes = requests.get(
        PDF_URL, timeout=10, headers={"User-Agent": "Mozilla/5.0"}
    ).content
    response = upload("test-report.pdf", pdf_bytes, "application/pdf")
    print(f"[pdf upload] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["file_type"] == "pdf"
    assert body["chunk_count"] > 0
    assert body["status"] == "stored"
    assert_stored_and_matches_db(body)


def run_valid_txt_upload_check():
    text = "Solar power converts sunlight into electricity using photovoltaic cells. " * 3
    response = upload("test-notes.txt", text.encode("utf-8"), "text/plain")
    print(f"[txt upload] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["file_type"] == "txt"
    assert body["chunk_count"] > 0
    assert_stored_and_matches_db(body)


def run_valid_docx_upload_check():
    document = DocxDocument()
    document.add_paragraph("Wind turbines convert kinetic energy from wind into electricity.")
    document.add_paragraph("Modern turbines automatically shut down above a safe wind speed.")
    buffer = io.BytesIO()
    document.save(buffer)
    response = upload(
        "test-memo.docx",
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    print(f"[docx upload] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["file_type"] == "docx"
    assert body["chunk_count"] > 0
    assert_stored_and_matches_db(body)


def run_empty_file_check():
    response = upload("test-empty.txt", b"", "text/plain")
    print(f"[empty file] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 400
    assert response.json()["detail"] == "Uploaded file is empty"
    assert get_document_row_by_filename("test-empty.txt") is None, (
        "No row should be written for a rejected empty upload"
    )


def run_unsupported_type_check():
    response = upload(
        "test-virus.exe", b"not a real executable but bytes anyway", "application/octet-stream"
    )
    print(f"[unsupported type] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]
    assert get_document_row_by_filename("test-virus.exe") is None


def run_multi_chunk_check():
    long_text = "Hydropower stores energy by pumping water uphill when demand is low. " * 40
    assert len(long_text) > CHUNK_SIZE * 2, "Fixture text must be long enough to force >1 chunk"
    response = upload("test-long.txt", long_text.encode("utf-8"), "text/plain")
    print(f"[multi chunk] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 200, response.text
    body = response.json()
    assert body["chunk_count"] > 1, "Expected a long document to produce multiple chunks"
    assert_stored_and_matches_db(body)


def run_corrupt_file_check():
    response = upload("test-corrupt.pdf", b"this is not a valid pdf file at all", "application/pdf")
    print(f"[corrupt pdf] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 422
    assert "corrupted or malformed" in response.json()["detail"]
    assert get_document_row_by_filename("test-corrupt.pdf") is None

    response = upload(
        "test-corrupt.docx",
        b"this is not a valid docx zip either",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    print(f"[corrupt docx] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 422
    assert get_document_row_by_filename("test-corrupt.docx") is None


def run_no_extractable_text_check():
    response = upload("test-blank.txt", b"   \n\n   ", "text/plain")
    print(f"[no extractable text] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 422
    assert "No extractable text" in response.json()["detail"]
    assert get_document_row_by_filename("test-blank.txt") is None


def run_file_too_large_check():
    oversized = b"a" * (MAX_FILE_SIZE_BYTES + 1)
    response = upload("test-toobig.txt", oversized, "text/plain")
    print(f"[file too large] status={response.status_code} body={response.json()}\n")
    assert response.status_code == 413
    assert get_document_row_by_filename("test-toobig.txt") is None


def run_missing_file_check():
    response = client.post("/documents/upload")  # no `file` part at all
    print(f"[missing file] status={response.status_code}\n")
    assert response.status_code == 422  # automatic, via FastAPI/Pydantic


def run():
    clear_previous_test_documents()
    run_valid_pdf_upload_check()
    run_valid_txt_upload_check()
    run_valid_docx_upload_check()
    run_empty_file_check()
    run_unsupported_type_check()
    run_multi_chunk_check()
    run_corrupt_file_check()
    run_no_extractable_text_check()
    run_file_too_large_check()
    run_missing_file_check()
    print(
        "All checks passed: /documents/upload validates, extracts, chunks, and "
        "durably stores PDF/TXT/DOCX uploads, and rejects every error case correctly."
    )


if __name__ == "__main__":
    run()
